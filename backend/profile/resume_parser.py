"""Resume parser — extracts text from PDF / DOCX / TXT for LLM ingestion.

Plain text extraction is good enough; the LLM downstream understands rough
formatting. We don't try to parse structure (headings, bullets, etc.) —
that's the LLM's job.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

# Lazy imports so the module loads even if optional deps aren't installed
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def parse_resume(path: str | Path) -> str:
    """Extract plain text from a resume file. Supports PDF, DOCX, TXT.

    Returns plain text with reasonable whitespace normalization.
    Raises FileNotFoundError if file doesn't exist.
    Raises ValueError if file format is unsupported.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Resume not found: {p}")

    suffix = p.suffix.lower()

    if suffix == ".pdf":
        if not HAS_PYPDF:
            raise ValueError(
                "pypdf not installed. Run: pip install pypdf"
            )
        return _parse_pdf(p)

    if suffix == ".docx":
        if not HAS_DOCX:
            raise ValueError(
                "python-docx not installed. Run: pip install python-docx"
            )
        return _parse_docx(p)

    if suffix in (".txt", ".md"):
        return _parse_text(p)

    raise ValueError(
        f"Unsupported resume format: {suffix}. "
        "Supported: .pdf, .docx, .txt, .md"
    )


def _parse_pdf(path: Path) -> str:
    """Extract text from PDF using pypdf."""
    reader = pypdf.PdfReader(str(path))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return _normalize_whitespace("\n\n".join(pages))


def _parse_docx(path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    doc = python_docx.Document(str(path))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Tables (some resumes use them for layout)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _normalize_whitespace("\n".join(parts))


def _parse_text(path: Path) -> str:
    """Read a plain text file."""
    return _normalize_whitespace(
        path.read_text(encoding="utf-8", errors="replace")
    )


def _normalize_whitespace(text: str) -> str:
    """Collapse repeated whitespace, normalize line breaks."""
    # Collapse 3+ newlines to 2 (preserves paragraph breaks)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Trim trailing whitespace per line
    lines = [line.rstrip() for line in text.split("\n")]
    text = "\n".join(lines)
    # Collapse runs of spaces/tabs (not newlines)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()
